"""
DOCX upload and text extraction module for AI Tutor application.
Handles DOCX files and extracts text using python-docx.
"""
import os
from typing import Tuple, Optional
import docx

class DOCXHandler:
    """Handles DOCX uploads and text extraction."""
    
    def __init__(self, upload_folder: str = "uploads/docx"):
        """
        Initialize the DOCX handler.
        
        Args:
            upload_folder: Directory to store uploaded DOCX files
        """
        self.upload_folder = upload_folder
        os.makedirs(upload_folder, exist_ok=True)
    
    def save_docx(self, docx_file, filename: str) -> str:
        """
        Save an uploaded DOCX file to disk.
        
        Args:
            docx_file: The uploaded DOCX file object
            filename: Name to save the file as
            
        Returns:
            Path to the saved DOCX file
        """
        file_path = os.path.join(self.upload_folder, filename)
        
        # Save the DOCX file
        with open(file_path, 'wb') as f:
            f.write(docx_file.read())
            
        return file_path
    
    def extract_text(self, docx_path: str) -> str:
        """
        Extract text from a DOCX file using python-docx.
        
        Args:
            docx_path: Path to the DOCX file
            
        Returns:
            Extracted text from the DOCX
        """
        try:
            doc = docx.Document(docx_path)
            
            # Extract text from paragraphs
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text)
                    full_text.append(" | ".join(row_text))
            
            return "\n".join(full_text)
        except Exception as e:
            return f"Error extracting text: {str(e)}"
    
    def process_docx(self, docx_file, filename: str) -> Tuple[str, Optional[str]]:
        """
        Process an uploaded DOCX: save it and extract text.
        
        Args:
            docx_file: The uploaded DOCX file object
            filename: Name to save the file as
            
        Returns:
            Tuple containing (file_path, extracted_text)
        """
        file_path = self.save_docx(docx_file, filename)
        extracted_text = self.extract_text(file_path)
        
        return file_path, extracted_text
